
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Cm
from urllib.parse import urljoin
import tkinter as tk
from tkinter import ttk, scrolledtext, messagebox, filedialog
import threading
import csv
import json
import os
import time
import logging
from concurrent.futures import ThreadPoolExecutor, as_completed

BASE_URL = "http://erogamescape.dyndns.org"
START_URL = "/~ap2/ero/toukei_kaiseki/toukei_avg.php?count=100&year=1900"

class App:
    def __init__(self, root):
        self.root = root
        self.root.title("游戏数据采集器")
        self.all_data = []
        self.config = self.load_config()
        self.is_running = False
        self.logger = self.setup_logger()
        self.create_widgets()

    def setup_logger(self):
        """设置日志记录"""
        logger = logging.getLogger('crawler')
        logger.setLevel(logging.INFO)
        fh = logging.FileHandler('crawler.log')
        fh.setLevel(logging.INFO)
        formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
        fh.setFormatter(formatter)
        logger.addHandler(fh)
        return logger

    def create_widgets(self):
        """创建界面组件"""
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)

        control_frame = ttk.Frame(main_frame)
        control_frame.grid(row=0, column=0, columnspan=2, pady=10, sticky=tk.W+tk.E)

        self.start_btn = ttk.Button(control_frame, text="开始采集", command=self.start_crawling)
        self.start_btn.pack(side=tk.LEFT, padx=5)
        self.stop_btn = ttk.Button(control_frame, text="停止", command=self.stop_crawling, state=tk.DISABLED)
        self.stop_btn.pack(side=tk.LEFT, padx=5)
        self.config_btn = ttk.Button(control_frame, text="配置", command=self.show_config)
        self.config_btn.pack(side=tk.LEFT, padx=5)
        self.preview_btn = ttk.Button(control_frame, text="预览数据", command=self.preview_data, state=tk.DISABLED)
        self.preview_btn.pack(side=tk.LEFT, padx=5)

        self.log_area = scrolledtext.ScrolledText(main_frame, width=80, height=20, wrap=tk.WORD)
        self.log_area.grid(row=1, column=0, columnspan=2, padx=10, pady=5, sticky=tk.W+tk.E+tk.N+tk.S)
        self.log_area.configure(state='disabled')

        progress_frame = ttk.Frame(main_frame)
        progress_frame.grid(row=2, column=0, columnspan=2, padx=10, pady=5, sticky=tk.W+tk.E)
        self.progress = ttk.Progressbar(progress_frame, orient=tk.HORIZONTAL, mode='determinate', length=300)
        self.progress.pack(side=tk.LEFT, fill=tk.X, expand=True)
        self.progress_label = ttk.Label(progress_frame, text="0%")
        self.progress_label.pack(side=tk.LEFT, padx=(10, 0))

        self.status_label = ttk.Label(main_frame, text="就绪")
        self.status_label.grid(row=3, column=0, columnspan=2, pady=5)

        main_frame.columnconfigure(0, weight=1)
        main_frame.rowconfigure(1, weight=1)

    def log(self, message):
        """记录日志到界面和文件"""
        self.logger.info(message)
        self.log_area.configure(state='normal')
        self.log_area.insert(tk.END, message + "\n")
        self.log_area.see(tk.END)
        self.log_area.configure(state='disabled')

    def update_status(self, text):
        """更新状态标签"""
        self.status_label.config(text=text)

    def start_crawling(self):
        """开始抓取"""
        self.is_running = True
        self.start_btn.config(state=tk.DISABLED)
        self.stop_btn.config(state=tk.NORMAL)
        self.config_btn.config(state=tk.DISABLED)
        self.preview_btn.config(state=tk.DISABLED)
        self.log_area.configure(state='normal')
        self.log_area.delete(1.0, tk.END)
        self.log_area.configure(state='disabled')
        self.all_data = []
        threading.Thread(target=self.main_process, daemon=True).start()

    def stop_crawling(self):
        """停止抓取"""
        self.is_running = False
        self.log("操作已中止")
        self.update_status("已中止")
        self.start_btn.config(state=tk.NORMAL)
        self.stop_btn.config(state=tk.DISABLED)
        self.config_btn.config(state=tk.NORMAL)
        self.preview_btn.config(state=tk.NORMAL)

    def main_process(self):
        """主抓取流程"""
        try:
            current_url = urljoin(BASE_URL, START_URL)
            total = 0
            urls = [current_url]

            while urls and self.is_running:
                with ThreadPoolExecutor(max_workers=5) as executor:
                    future_to_url = {executor.submit(self.fetch_page, url): url for url in urls}
                    urls = []
                    for future in as_completed(future_to_url):
                        soup = future.result()
                        if soup:
                            page_data = self.parse_table(soup)
                            filtered_data = self.filter_data(page_data)
                            self.all_data.extend(filtered_data)
                            total += len(filtered_data)
                            self.log(f"已获取 {len(filtered_data)} 条数据，总计 {total} 条")

                            next_url = self.get_next_page(soup)
                            if next_url and len(self.all_data) < self.config['max_items']:
                                urls.append(next_url)

                            self.update_progress((len(self.all_data) / self.config['max_items']) * 100)

                        if len(self.all_data) >= self.config['max_items']:
                            self.log("已达到最大采集数量，停止采集")
                            break

            if self.is_running:
                self.save_data()
                messagebox.showinfo("完成", f"数据采集完成，共 {total} 条记录")
                self.update_status("完成")

        except Exception as e:
            messagebox.showerror("错误", str(e))
        finally:
            self.start_btn.config(state=tk.NORMAL)
            self.stop_btn.config(state=tk.DISABLED)
            self.config_btn.config(state=tk.NORMAL)
            self.preview_btn.config(state=tk.NORMAL)
            self.progress['value'] = 0
            self.update_progress(0)

    def update_progress(self, percentage):
        """更新进度条"""
        self.progress['value'] = percentage
        self.progress_label.config(text=f"{percentage:.1f}%")

    def save_data(self):
        """保存数据"""
        file_types = [("Word Documents", "*.docx"), ("CSV Files", "*.csv"), ("Excel Files", "*.xlsx")]
        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=file_types, title="保存文件")
        if file_path:
            if file_path.endswith('.docx'):
                self.save_to_word(self.all_data, file_path)
            elif file_path.endswith('.csv'):
                self.save_to_csv(self.all_data, file_path)
            elif file_path.endswith('.xlsx'):
                self.save_to_excel(self.all_data, file_path)
            self.log(f"数据已保存至 {file_path}")

    def filter_data(self, data):
        """过滤数据"""
        return [row for row in data if float(row[2]) >= self.config['min_avg_score']]

    def show_config(self):
        """显示配置窗口"""
        config_window = tk.Toplevel(self.root)
        config_window.title("配置")

        ttk.Label(config_window, text="每页抓取数量:").grid(row=0, column=0, padx=5, pady=5)
        count_entry = ttk.Entry(config_window)
        count_entry.insert(0, str(self.config['count']))
        count_entry.grid(row=0, column=1, padx=5, pady=5)

        ttk.Label(config_window, text="起始年份:").grid(row=1, column=0, padx=5, pady=5)
        year_entry = ttk.Entry(config_window)
        year_entry.insert(0, str(self.config['year']))
        year_entry.grid(row=1, column=1, padx=5, pady=5)

        ttk.Label(config_window, text="最大采集数量:").grid(row=2, column=0, padx=5, pady=5)
        max_items_entry = ttk.Entry(config_window)
        max_items_entry.insert(0, str(self.config['max_items']))
        max_items_entry.grid(row=2, column=1, padx=5, pady=5)

        ttk.Label(config_window, text="最低平均分:").grid(row=3, column=0, padx=5, pady=5)
        min_score_entry = ttk.Entry(config_window)
        min_score_entry.insert(0, str(self.config['min_avg_score']))
        min_score_entry.grid(row=3, column=1, padx=5, pady=5)

        def save_config():
            self.config['count'] = int(count_entry.get())
            self.config['year'] = int(year_entry.get())
            self.config['max_items'] = int(max_items_entry.get())
            self.config['min_avg_score'] = float(min_score_entry.get())
            self.save_config()
            config_window.destroy()

        ttk.Button(config_window, text="保存", command=save_config).grid(row=4, column=0, columnspan=2, pady=10)

    def load_config(self):
        """加载配置"""
        default_config = {'count': 100, 'year': 1900, 'max_items': 1000, 'min_avg_score': 0}
        if os.path.exists('config.json'):
            with open('config.json', 'r') as f:
                return json.load(f)
        return default_config

    def save_config(self):
        """保存配置"""
        with open('config.json', 'w') as f:
            json.dump(self.config, f)

    def preview_data(self):
        """预览数据"""
        preview_window = tk.Toplevel(self.root)
        preview_window.title("数据预览")
        preview_text = scrolledtext.ScrolledText(preview_window, width=80, height=20)
        preview_text.pack(padx=10, pady=10)
        headers = ['游戏名', '品牌名', '平均值', '中位数', '标准偏差', '数据数']
        preview_text.insert(tk.END, '\t'.join(headers) + '\n')
        for row in self.all_data[:100]:
            preview_text.insert(tk.END, '\t'.join(row) + '\n')
        preview_text.configure(state='disabled')

    def fetch_page(self, url, retries=3):
        """抓取页面，带重试机制"""
        for attempt in range(retries):
            try:
                response = requests.get(url, timeout=10)
                response.encoding = 'utf-8'
                return BeautifulSoup(response.text, 'html.parser')  # 可替换为 'lxml' 提高性能
            except Exception as e:
                self.log(f"请求失败: {e}，重试 {attempt+1}/{retries}")
                time.sleep(2)  # 延迟避免频繁请求
        return None

    def parse_table(self, soup):
        """解析表格"""
        table = soup.find('table')
        rows = []
        for tr in table.find_all('tr')[1:]:
            cols = [td.get_text(strip=True) for td in tr.find_all(['th', 'td'])]
            if len(cols) == 6:
                rows.append(cols)
        return rows

    def get_next_page(self, soup):
        """获取下一页链接"""
        link = soup.find('a', string='次の100件を見る')
        return urljoin(BASE_URL, link['href']) if link else None

    def save_to_word(self, data, filename):
        """保存为 Word"""
        doc = Document()
        doc.add_heading('游戏统计表', 0)
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        widths = (Cm(5), Cm(4), Cm(2), Cm(2), Cm(2), Cm(2))
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width
        hdr_cells = table.rows[0].cells
        headers = ['游戏名', '品牌名', '平均值', '中位数', '标准偏差', '数据数']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
        for row in data:
            row_cells = table.add_row().cells
            for i in range(6):
                row_cells[i].text = row[i]
        doc.save(filename)

    def save_to_csv(self, data, filename):
        """保存为 CSV"""
        with open(filename, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerow(['游戏名', '品牌名', '平均值', '中位数', '标准偏差', '数据数'])
            writer.writerows(data)

    def save_to_excel(self, data, filename):
        """保存为 Excel"""
        import openpyxl
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.append(['游戏名', '品牌名', '平均值', '中位数', '标准偏差', '数据数'])
        for row in data:
            ws.append(row)
        wb.save(filename)

if __name__ == "__main__":
    root = tk.Tk()
    app = App(root)
    root.mainloop()